from __future__ import annotations

import argparse
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


MAX_IMAGE_WIDTH_CM = 15.5


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for style_name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_runs(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    if node.name == "img":
        return

    run_style = {
        "bold": node.name in {"strong", "b"},
        "italic": node.name in {"em", "i"},
        "underline": node.name == "u",
    }

    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = run_style["bold"]
                run.italic = run_style["italic"]
                run.underline = run_style["underline"]
        else:
            add_runs(paragraph, child)


def add_heading(document: Document, element: Tag, is_first_heading: bool) -> None:
    level = min(int(element.name[1]), 3)
    text = element.get_text(" ", strip=True)
    if is_first_heading:
        paragraph = document.add_paragraph(text, style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, element: Tag) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    for child in element.children:
        add_runs(paragraph, child)
    if not paragraph.text.strip():
        paragraph.clear()


def add_list(document: Document, element: Tag) -> None:
    style = "List Number" if element.name == "ol" else "List Bullet"
    for item in element.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        for child in item.children:
            if isinstance(child, Tag) and child.name in {"ol", "ul"}:
                add_list(document, child)
            else:
                add_runs(paragraph, child)


def add_table(document: Document, element: Tag) -> None:
    rows = element.find_all("tr")
    if not rows:
        return

    first_row = rows[0].find_all(["th", "td"], recursive=False)
    table = document.add_table(rows=0, cols=len(first_row))
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        cells = row.find_all(["th", "td"], recursive=False)
        table_row = table.add_row().cells
        for cell_index, cell in enumerate(cells):
            table_cell = table_row[cell_index]
            table_cell.text = cell.get_text(" ", strip=True)
            if row_index == 0:
                set_cell_shading(table_cell, "D9EAF7")
                for paragraph in table_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_image(document: Document, element: Tag, source_path: Path) -> None:
    image_path = (source_path.parent / element["src"]).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[Missing image: {element['src']}]")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(MAX_IMAGE_WIDTH_CM))


def convert_markdown(source_path: Path, output_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8")
    html = markdown.markdown(text, extensions=["tables"])
    soup = BeautifulSoup(html, "html.parser")

    document = Document()
    configure_document(document)

    first_heading = True
    for element in soup.contents:
        if not isinstance(element, Tag):
            continue

        if element.name in {"h1", "h2", "h3"}:
            add_heading(document, element, first_heading)
            first_heading = False
        elif element.name == "p":
            image = element.find("img", recursive=False)
            if image is not None:
                add_image(document, image, source_path)
            else:
                add_paragraph(document, element)
        elif element.name in {"ol", "ul"}:
            add_list(document, element)
        elif element.name == "table":
            add_table(document, element)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert Markdown report to DOCX.")
    parser.add_argument("source", type=Path, help="Path to the markdown file.")
    parser.add_argument("output", type=Path, nargs="?", help="Output DOCX path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source_path = args.source.resolve()
    output_path = args.output.resolve() if args.output else source_path.with_suffix(".docx")
    convert_markdown(source_path, output_path)


if __name__ == "__main__":
    main()